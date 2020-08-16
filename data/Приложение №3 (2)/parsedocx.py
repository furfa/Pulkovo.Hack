import os
from pprint import pprint
import subprocess
import docx
import json
path = "./"
docs = os.listdir(path)

def check_cell(cell):
	cell = cell.strip()
	return (cell != "-" and cell != "–" and cell != "")

for filename in docs:
    if filename.endswith(".doc") or filename.endswith(".DOC"):
        subprocess.call(["soffice", "--headless", "--convert-to", "docx", path+filename])

for filename in docs:
    if filename.endswith(".doc") or filename.endswith(".DOC"):
        subprocess.call(["rm", path+filename])

docs = os.listdir(path)

docs.remove("parsedocx.py")
tmp_files = [doc for doc in docs if doc.endswith("#")]
for tmp_file in tmp_files:
	docs.remove(tmp_file)

res = {}
for p, file in enumerate(docs):
	
	document = docx.Document(path+file)
	course = document.sections[0].header.tables[0].rows[0].cells[1].paragraphs[2:]
	course = " ".join([para.text.strip().replace("\n", " ") for para in course])
	res[course] = {}
	res[course]["file"] = file
	for table in document.tables[:1]:
		c = 0
		section = ""
		remove_first = False
		if (table.rows[0].cells[0].paragraphs[0].text.strip() == "№ п/п" or
			table.rows[0].cells[0].paragraphs[0].text.strip() == "№п/п" or
		    table.rows[0].cells[0].paragraphs[0].text.strip() == "№"):
			remove_first = True
		for row in table.rows:
			cells = [cell.paragraphs[0].text.replace("\xa0", " ") for cell in row.cells]
			if remove_first:
				cells = cells[1:]
			if cells[0] == cells[1]:
				cells.pop(0)
			
			#print(cells)
			tmp = cells[0].split()
			cells[0] = []
			cells[0].append(" ".join(tmp[:2]))
			cells[0][0] = cells[0][0].replace("№", "").replace("\n", " ").replace("  ", " ")
			cells[0].append(" ".join(tmp[2:]))
			cells[0][1] = cells[0][1].replace("№", "").replace("\n", " ").replace("  ", " ")
			
			if cells[0][0][:6] == "Раздел":
				section, name = cells[0][0], cells[0][1]
				if section.endswith("."):
					section = section[:-1]
				if cells[1] == "–" or cells[1] == "":
					name = cells[4]
					cells[1] = "0"
				res[course][section] = {
					"name": name, 
					"time": float(cells[1].replace(",", ".")),
					"topics": []
				}
				c = 1
			if c == 1 and cells[0][0][:4] == "Тема":
				if cells[0][0].endswith("."):
					cells[0][0] = cells[0][0][:-1]
				if cells[1] == "–" or cells[1] == "":
					cells[1] = "0"
				res[course][section]["topics"].append(
					(
						cells[0][1], # topic name
						float(cells[2].replace(",", ".") if check_cell(cells[2]) else "0.0"),
						float(cells[3].replace(",", ".") if check_cell(cells[3]) else "0.0") # duration
					)
				)
	print(int((p+1)/len(docs)*1000)/10)

with open("./../parsed_docx.json", "w") as fp:
	json.dump(res, fp)